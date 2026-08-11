from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument


class TextExtractionError(Exception):
    pass


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF. Falls back to OCR-friendly
    behavior is not automatic here; pytesseract can be wired in for scanned PDFs
    if a page yields no text."""
    text_parts = []
    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(page_text)
    except Exception as e:
        raise TextExtractionError(f"Failed to extract PDF text: {e}")

    return "\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = DocxDocument(file_path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]

        # Include table cell text as well
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)

        return "\n".join(parts)
    except Exception as e:
        raise TextExtractionError(f"Failed to extract DOCX text: {e}")


def extract_text_from_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        raise TextExtractionError(f"Failed to extract TXT text: {e}")


def extract_text(file_path: str, file_type: str) -> str:
    """Dispatch extraction based on file extension/type."""
    file_type = file_type.lower().lstrip(".")

    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type == "docx":
        return extract_text_from_docx(file_path)
    elif file_type == "txt":
        return extract_text_from_txt(file_path)
    else:
        raise TextExtractionError(f"Unsupported file type: {file_type}")


def get_file_type(filename: str) -> str:
    ext = Path(filename).suffix.lower().lstrip(".")
    if ext not in ("pdf", "docx", "txt"):
        raise TextExtractionError(f"Unsupported file extension: {ext}")
    return ext
